"""
A collection of math question generators.

parameters:
- page_break: bool, True if this page requires a page break.
- start: bool, True if this page is the start.
- length: int; = 1, is changed if there is more than 1 page ahead of this page.

Call seprate functions for seperate pages.
"""

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
import random
import sys
import os

path = sys.path[0] + "\\printing documents"

def additionGenerator20(page_break: bool, start: bool, length = 1):
    """Rewrites in additionIn20.docx, includes equations of numbers from 2-9."""

    addition_printing = Document(path + "\\additionIn20.docx")
    
    if start:
            try:
                for para in addition_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = addition_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    addition_printing.add_paragraph().add_run("date:                               time:                               ___/64")
    
    table = addition_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(2, 10):
        all_numbers[i] = []
        for j in range(2, 10):
            all_numbers[i].append(j)
            counter += 1
    
    for i in range(0, counter):
        first_num = random.randint(2, 9)
        second_num = random.randint(2, 9)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(2, 9)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(2, 9)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "+" + str(second_num) + "="
    
    if page_break:
        addition_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    addition_printing.save(path + "\\additionIn20.docx")

def subtractionGenerator20(page_break: bool, start: bool, length = 1):
    """Rewrites in subtractionIn20.docx, includes equations of subtrahend from 4-18, minuhend from 2-9, 
    excluding results of more than 10 and 1."""
    
    subtraction_printing = Document(path + "\\subtraction_printing.docx")
    
    if start:
            try:
                for para in subtraction_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = subtraction_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    subtraction_printing.add_paragraph().add_run("date:                               time:                               ___/64")

    table = subtraction_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(4, 19):
        all_numbers[i] = []
        for j in range(2, 10):
            all_numbers[i].append(j)
            counter += 1

    for i in range(4, 19):
        for j in range(2,10):
            if (i-j)/10 >= 1 or (i-j) < 2:
                all_numbers[i].remove(j)
                counter -= 1

    for i in range(0, counter):
        first_num = random.randint(4, 18)
        second_num = random.randint(2, 9)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(4, 18)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(2, 9)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "-" + str(second_num) + "="
        
    if page_break:
        subtraction_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    subtraction_printing.save(path + "\\subtraction_printing.docx")

def additionGenerator50(page_break: bool, start: bool, length = 1):
    """Rewrites in additionIn50.docx, includes equations of numbers from 12-49, excluding equations
    containg the same numbers."""

    subtraction_printing = Document(path + "\\additionIn50.docx")

    start_num = 12
    end_num = 49
    
    if start:
            try:
                for para in subtraction_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = subtraction_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    subtraction_printing.add_paragraph().add_run("date:                               time:                               ___/64")

    table = subtraction_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(start_num, end_num+1):
        all_numbers[i] = []
        for j in range(start_num, end_num+1):
            all_numbers[i].append(j)
            counter += 1

    for i in range(start_num, end_num+1):
        for j in range(start_num, end_num+1):
            if (j%10 <= 1) or (i%10 <= 1) or (i in all_numbers[j]):
                all_numbers[i].remove(j)
                counter -= 1

    for i in range(0, 64):
        first_num = random.randint(start_num, end_num)
        second_num = random.randint(start_num, end_num)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(start_num, end_num)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(start_num, end_num)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "+" + str(second_num) + "="
        
    if page_break:
        subtraction_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    subtraction_printing.save(path + "\\additionIn50.docx")

def subtractionGenerator50(page_break: bool, start: bool, length = 1):
    """Rewrites in subtractionIn50.docx, includes equations of subtrahend from 24-99, minuhend from 12-87, 
    excluding numbers divisible by 10, single digit answers and answers with 1 as their last digit."""

    subtraction_printing = Document(path + "\\subtractionIn50.docx")

    start_num = 24
    end_num = 99
    smallest_subtractable = 12
    
    if start:
            try:
                for para in subtraction_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = subtraction_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    subtraction_printing.add_paragraph().add_run("date:                               time:                               ___/64")

    table = subtraction_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(start_num, end_num+1):
        all_numbers[i] = []
        for j in range(start_num-smallest_subtractable, end_num-smallest_subtractable+1):
            all_numbers[i].append(j)
            counter += 1

    for i in range(start_num, end_num+1):
        for j in range(start_num-smallest_subtractable, end_num-smallest_subtractable+1):
            if ((i-j)%10 <= 1) or ((i-j) < 10) or (j%10 == 0) or (i%10 == 0):
                all_numbers[i].remove(j)
                counter -= 1

    for i in range(0, 64):
        first_num = random.randint(start_num, end_num)
        second_num = random.randint(start_num-smallest_subtractable, end_num-smallest_subtractable)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(start_num, end_num)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(start_num-smallest_subtractable, end_num-smallest_subtractable)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "-" + str(second_num) + "="
        
    if page_break:
        subtraction_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    subtraction_printing.save(path + "\\subtractionIn50.docx")