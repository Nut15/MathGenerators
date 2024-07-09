from docx import Document
from docx.shared import Cm
import random

def additionGenerator():
    
    addition_printing = Document("addition_printing.docx")
    
    try:
        table = addition_printing.tables[0]
        table._element.getparent().remove(table._element)
    except:
        addition_printing.paragraphs[0].add_run("date:                               time:                               ___/64")
        pass
    
    table = addition_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(2, 10):
        all_numbers[i] = []
        for j in range(2, 10):
            all_numbers[i].append(j)
            counter += 1
    
    for i in range(0, counter):
        first_num = random.randint(2, 9)
        second_num = random.randint(2, 9)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(2, 9)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(2, 9)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "+" + str(second_num) + "="

    
    addition_printing.save("addition_printing.docx")