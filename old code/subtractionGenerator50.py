from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
import random
import sys

def subtractionGenerator(page_break, start, length = 1):

    path = sys.path[0].replace("old code","printing documents")
    
    subtraction_printing = Document(path + "\\subtractionIn50.docx")

    start_num = 24
    end_num = 99
    smallest_subtractable = 12
    
    if start:
            try:
                for para in subtraction_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = subtraction_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    subtraction_printing.add_paragraph().add_run("date:                               time:                               ___/64")

    table = subtraction_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(start_num, end_num+1):
        all_numbers[i] = []
        for j in range(start_num-smallest_subtractable, end_num-smallest_subtractable+1):
            all_numbers[i].append(j)
            counter += 1

    for i in range(start_num, end_num+1):
        for j in range(start_num-smallest_subtractable, end_num-smallest_subtractable+1):
            if ((i-j)%10 <= 1) or ((i-j) < 10) or (j%10 == 0) or (i%10 == 0):
                all_numbers[i].remove(j)
                counter -= 1
    print(all_numbers)

    print(counter)
    for i in range(0, 64):
        first_num = random.randint(start_num, end_num)
        second_num = random.randint(start_num-smallest_subtractable, end_num-smallest_subtractable)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(start_num, end_num)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(start_num-smallest_subtractable, end_num-smallest_subtractable)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                print(str(first_num) + "-" + str(second_num) + "=")
                cell.text = str(first_num) + "-" + str(second_num) + "="
        
    if page_break:
        subtraction_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    subtraction_printing.save(path + "\\subtractionIn50.docx")

subtractionGenerator(False, True)