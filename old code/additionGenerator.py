from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
import random
import sys

def additionGenerator(page_break, start, length = 1):

    path = sys.path[0].replace("old code","printing documents")

    addition_printing = Document(path + "\\additionIn20.docx")
    
    if start:
            try:
                for para in addition_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = addition_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    addition_printing.add_paragraph().add_run("date:                               time:                               ___/64")
    
    table = addition_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(2, 10):
        all_numbers[i] = []
        for j in range(2, 10):
            all_numbers[i].append(j)
            counter += 1
    
    for i in range(0, counter):
        first_num = random.randint(2, 9)
        second_num = random.randint(2, 9)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(2, 9)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(2, 9)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "+" + str(second_num) + "="
    
    if page_break:
        addition_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    addition_printing.save(path + "\\additionIn20.docx")