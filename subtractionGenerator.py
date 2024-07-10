from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
import random

def subtractionGenerator(page_break, start, length = 1):
    
    subtraction_printing = Document("subtraction_printing.docx")
    
    if start:
            try:
                for para in subtraction_printing.paragraphs:
                    para._element.getparent().remove(para._element)
                for i in range(0, length):
                    table = subtraction_printing.tables[0]
                    table._element.getparent().remove(table._element)
            except:
                pass
    subtraction_printing.add_paragraph().add_run("date:                               time:                               ___/64")

    table = subtraction_printing.add_table(rows=1, cols=4)
    table.autofit = True
    
    all_numbers = {}
    counter = 0
    for i in range(4, 19):
        all_numbers[i] = []
        for j in range(2, 10):
            all_numbers[i].append(j)
            counter += 1

    for i in range(4, 19):
        for j in range(2,10):
            if (i-j)/10 >= 1 or (i-j) < 2:
                all_numbers[i].remove(j)
                counter -= 1

    for i in range(0, counter):
        first_num = random.randint(4, 18)
        second_num = random.randint(2, 9)
        while len(all_numbers[first_num]) == 0:
            first_num = random.randint(4, 18)
        else:
            while second_num not in all_numbers[first_num]:
                second_num = random.randint(2, 9)
            else:
                all_numbers[first_num].remove(second_num)
                if i%4 == 0:
                    table.add_row()
                cell = table.cell(int(i/4)+1, int(i%4))
                cell.width = Cm(29.7/4)
                cell.text = str(first_num) + "-" + str(second_num) + "="
        
    if page_break:
        subtraction_printing.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    subtraction_printing.save("subtraction_printing.docx")
